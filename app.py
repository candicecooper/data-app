import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime, date, time, timedelta
import uuid
import random
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import tempfile
import os

# =========================================
# CONFIG + CONSTANTS
# =========================================

st.set_page_config(
    page_title="CLC Behaviour Support – SANDBOX",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# =========================================
# REFINED MODERN STYLING - SLICK & READABLE
# =========================================

st.markdown("""
<style>
    /* Import modern font */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Global styling */
    * {
        font-family: 'Inter', sans-serif;
    }
    
    /* Refined dark background - much more sophisticated */
    .stApp {
        background: linear-gradient(135deg, #1a1d29 0%, #2d3748 50%, #1a202c 100%);
    }
    
    /* Custom buttons - refined navy/teal */
    .stButton>button {
        background: linear-gradient(135deg, #2c5282 0%, #2b6cb0 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.6rem 1.5rem !important;
        font-weight: 500 !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.3) !important;
    }
    
    .stButton>button:hover {
        background: linear-gradient(135deg, #2b6cb0 0%, #3182ce 100%) !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.4) !important;
    }
    
    /* Primary button - refined teal */
    button[kind="primary"] {
        background: linear-gradient(135deg, #0d9488 0%, #14b8a6 100%) !important;
        box-shadow: 0 2px 8px rgba(13, 148, 136, 0.4) !important;
    }
    
    button[kind="primary"]:hover {
        background: linear-gradient(135deg, #14b8a6 0%, #2dd4bf 100%) !important;
        box-shadow: 0 4px 12px rgba(13, 148, 136, 0.5) !important;
    }
    
    /* Containers with excellent readability */
    [data-testid="stVerticalBlock"] > div[style*="border"] {
        background: rgba(255, 255, 255, 0.98) !important;
        border-radius: 12px !important;
        padding: 2rem !important;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
    }
    
    /* Metrics - refined colors */
    [data-testid="stMetricValue"] {
        font-size: 2rem !important;
        font-weight: 700 !important;
        color: #2c5282 !important;
    }
    
    [data-testid="stMetricLabel"] {
        color: #4a5568 !important;
        font-weight: 500 !important;
    }
    
    /* Input fields - clean and readable */
    .stTextInput>div>div>input,
    .stSelectbox>div>div>select,
    .stTextArea>div>div>textarea,
    .stDateInput>div>div>input,
    .stTimeInput>div>div>input,
    .stNumberInput>div>div>input {
        border-radius: 8px !important;
        border: 2px solid #cbd5e0 !important;
        background: white !important;
        color: #2d3748 !important;
        transition: all 0.3s ease !important;
    }
    
    .stTextInput>div>div>input:focus,
    .stSelectbox>div>div>select:focus,
    .stTextArea>div>div>textarea:focus {
        border-color: #2c5282 !important;
        box-shadow: 0 0 0 3px rgba(44, 82, 130, 0.1) !important;
    }
    
    /* Headers - white text with shadow for contrast */
    h1, h2, h3 {
        color: white !important;
        font-weight: 700 !important;
        text-shadow: 0 2px 4px rgba(0, 0, 0, 0.3) !important;
    }
    
    /* Subheaders in containers - dark text */
    [data-testid="stVerticalBlock"] h1,
    [data-testid="stVerticalBlock"] h2,
    [data-testid="stVerticalBlock"] h3,
    [data-testid="stVerticalBlock"] h4 {
        color: #2d3748 !important;
        text-shadow: none !important;
    }
    
    /* Expanders - refined dark glass */
    .streamlit-expanderHeader {
        background: rgba(255, 255, 255, 0.1) !important;
        border-radius: 8px !important;
        backdrop-filter: blur(10px) !important;
        color: white !important;
    }
    
    /* Success/info boxes - better contrast */
    .stSuccess {
        background: rgba(236, 253, 245, 0.98) !important;
        border-radius: 8px !important;
        border-left: 4px solid #059669 !important;
        color: #065f46 !important;
    }
    
    .stInfo {
        background: rgba(239, 246, 255, 0.98) !important;
        border-radius: 8px !important;
        border-left: 4px solid #2563eb !important;
        color: #1e40af !important;
    }
    
    .stWarning {
        background: rgba(254, 252, 232, 0.98) !important;
        border-radius: 8px !important;
        border-left: 4px solid #d97706 !important;
        color: #92400e !important;
    }
    
    .stError {
        background: rgba(254, 242, 242, 0.98) !important;
        border-radius: 8px !important;
        border-left: 4px solid #dc2626 !important;
        color: #991b1b !important;
    }
    
    /* Download buttons - refined green */
    .stDownloadButton>button {
        background: linear-gradient(135deg, #059669 0%, #10b981 100%) !important;
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
        box-shadow: 0 2px 8px rgba(5, 150, 105, 0.3) !important;
    }
    
    .stDownloadButton>button:hover {
        background: linear-gradient(135deg, #10b981 0%, #34d399 100%) !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 12px rgba(5, 150, 105, 0.4) !important;
    }
    
    /* Tabs - refined styling */
    .stTabs [data-baseweb="tab-list"] {
        background: rgba(255, 255, 255, 0.1) !important;
        border-radius: 8px !important;
        padding: 0.5rem !important;
    }
    
    .stTabs [data-baseweb="tab"] {
        color: rgba(255, 255, 255, 0.7) !important;
        border-radius: 6px !important;
        font-weight: 500 !important;
    }
    
    .stTabs [aria-selected="true"] {
        background: rgba(255, 255, 255, 0.2) !important;
        color: white !important;
    }
    
    /* Captions - better readability */
    .stCaption {
        color: rgba(255, 255, 255, 0.8) !important;
    }
    
    [data-testid="stVerticalBlock"] .stCaption {
        color: #718096 !important;
    }
    
    /* Sliders */
    .stSlider [data-baseweb="slider"] {
        background: rgba(255, 255, 255, 0.1) !important;
    }
    
    /* Checkbox and Radio */
    .stCheckbox, .stRadio {
        color: white !important;
    }
    
    [data-testid="stVerticalBlock"] .stCheckbox label,
    [data-testid="stVerticalBlock"] .stRadio label {
        color: #2d3748 !important;
    }
    
    /* Code blocks */
    .stCodeBlock {
        background: rgba(255, 255, 255, 0.95) !important;
        border-radius: 8px !important;
    }
    
    /* Dataframe */
    .stDataFrame {
        background: white !important;
        border-radius: 8px !important;
    }
</style>
""", unsafe_allow_html=True)

# --- Refined Sandbox banner ---
st.markdown("""
<div style='background: rgba(255, 255, 255, 0.95); 
            padding: 1.5rem; 
            border-radius: 12px; 
            margin-bottom: 2rem;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2);
            border-left: 5px solid #2c5282;'>
    <h3 style='color: #2c5282; margin: 0; font-size: 1.2rem; text-shadow: none;'>
        🎭 SANDBOX MODE
    </h3>
    <p style='color: #4a5568; margin: 0.5rem 0 0 0;'>
        This demonstration uses synthetic data only. No real student information is included.
    </p>
</div>
""", unsafe_allow_html=True)

# --- Mock Staff (with emails AND passwords) ---
MOCK_STAFF = [
    {"id": "s1", "name": "Emily Jones", "role": "JP", "email": "emily.jones@example.com", "password": "demo123"},
    {"id": "s2", "name": "Daniel Lee", "role": "PY", "email": "daniel.lee@example.com", "password": "demo123"},
    {"id": "s3", "name": "Sarah Chen", "role": "SY", "email": "sarah.chen@example.com", "password": "demo123"},
    {"id": "s4", "name": "Admin User", "role": "ADM", "email": "admin.user@example.com", "password": "admin123"},
    {"id": "s5", "name": "Michael Torres", "role": "JP", "email": "michael.torres@example.com", "password": "demo123"},
    {"id": "s6", "name": "Jessica Williams", "role": "PY", "email": "jessica.williams@example.com", "password": "demo123"},
]

# --- EXPANDED Mock Students (3 per program = 9 total) ---
MOCK_STUDENTS = [
    # JP - 3 students
    {"id": "stu_jp1", "name": "Emma T.", "grade": "R", "dob": "2018-05-30", "edid": "ED12348", "program": "JP"},
    {"id": "stu_jp2", "name": "Oliver S.", "grade": "Y1", "dob": "2017-09-12", "edid": "ED12349", "program": "JP"},
    {"id": "stu_jp3", "name": "Sophie M.", "grade": "Y2", "dob": "2016-03-20", "edid": "ED12350", "program": "JP"},
    
    # PY - 3 students  
    {"id": "stu_py1", "name": "Liam C.", "grade": "Y3", "dob": "2015-06-15", "edid": "ED12351", "program": "PY"},
    {"id": "stu_py2", "name": "Ava R.", "grade": "Y4", "dob": "2014-11-08", "edid": "ED12352", "program": "PY"},
    {"id": "stu_py3", "name": "Noah B.", "grade": "Y6", "dob": "2012-02-28", "edid": "ED12353", "program": "PY"},
    
    # SY - 3 students
    {"id": "stu_sy1", "name": "Isabella G.", "grade": "Y7", "dob": "2011-04-17", "edid": "ED12354", "program": "SY"},
    {"id": "stu_sy2", "name": "Ethan D.", "grade": "Y9", "dob": "2009-12-03", "edid": "ED12355", "program": "SY"},
    {"id": "stu_sy3", "name": "Mia A.", "grade": "Y11", "dob": "2007-08-20", "edid": "ED12356", "program": "SY"},
]

PROGRAM_NAMES = {"JP": "Junior Primary", "PY": "Primary Years", "SY": "Senior Years"}

SUPPORT_TYPES = [
    "1:1 Individual Support",
    "Independent",
    "Small Group",
    "Large Group",
]

BEHAVIOUR_TYPES = [
    "Verbal Refusal",
    "Elopement",
    "Property Destruction",
    "Aggression (Peer)",
    "Aggression (Adult)",
    "Self-Harm",
    "Verbal Aggression",
    "Other",
]

ANTECEDENTS = [
    "Requested to transition activity",
    "Given instruction / demand (Academic)",
    "Peer conflict / teasing",
    "Staff attention shifted away",
    "Unstructured free time (Recess/Lunch)",
    "Sensory overload (noise / lights)",
    "Access to preferred item denied",
    "Change in routine or expectation",
    "Difficult task presented",
]

INTERVENTIONS = [
    "Used calm tone and supportive stance (CPI)",
    "Offered a break / time away",
    "Reduced task demand / chunked task",
    "Provided choices",
    "Removed audience / peers",
    "Used visual supports",
    "Co-regulated with breathing / grounding",
    "Prompted use of coping skill",
    "Redirection to preferred activity",
]

LOCATIONS = [
    "JP Classroom",
    "JP Spill Out",
    "PY Classroom",
    "PY Spill Out",
    "SY Classroom",
    "SY Spill Out",
    "Student Kitchen",
    "Admin",
    "Gate",
    "Library",
    "Playground",
    "Yard",
    "Toilets",
    "Excursion",
    "Swimming",
]

VALID_PAGES = [
    "login",
    "landing",
    "program_students",
    "incident_log",
    "critical_incident",
    "student_analysis",
    "program_overview",
]

# =========================================
# HELPER FUNCTIONS
# =========================================


def init_state():
    """Initialise all session_state keys used in the app."""
    ss = st.session_state
    if "logged_in" not in ss:
        ss.logged_in = False
    if "current_user" not in ss:
        ss.current_user = None
    if "current_page" not in ss:
        ss.current_page = "login"
    if "students" not in ss:
        ss.students = MOCK_STUDENTS
    if "staff" not in ss:
        ss.staff = MOCK_STAFF
    if "incidents" not in ss:
        ss.incidents = generate_mock_incidents(70)
    if "critical_incidents" not in ss:
        ss.critical_incidents = []
    if "selected_program" not in ss:
        ss.selected_program = "JP"
    if "selected_student_id" not in ss:
        ss.selected_student_id = None
    if "current_incident_id" not in ss:
        ss.current_incident_id = None
    if "abch_rows" not in ss:
        ss.abch_rows = []


def login_user(email: str, password: str) -> bool:
    """
    Login with email AND password verification.
    """
    email = (email or "").strip().lower()
    password = (password or "").strip()
    
    if not email or not password:
        return False

    # Check against mock staff
    for staff in st.session_state.staff:
        if staff.get("email", "").lower() == email and staff.get("password", "") == password:
            st.session_state.logged_in = True
            st.session_state.current_user = staff
            st.session_state.current_page = "landing"
            return True
    
    # For demo purposes, also allow "demo" as universal password
    if password == "demo":
        for staff in st.session_state.staff:
            if staff.get("email", "").lower() == email:
                st.session_state.logged_in = True
                st.session_state.current_user = staff
                st.session_state.current_page = "landing"
                return True
        
        # Fallback demo user if email not found but password is "demo"
        demo = {"id": "demo_staff", "name": "Demo User", "role": "JP", "email": email}
        st.session_state.logged_in = True
        st.session_state.current_user = demo
        st.session_state.current_page = "landing"
        return True

    return False


def go_to(page: str, **kwargs):
    """Simple page navigation helper."""
    if page not in VALID_PAGES:
        st.error(f"Unknown page: {page}")
        return
    st.session_state.current_page = page
    for k, v in kwargs.items():
        setattr(st.session_state, k, v)
    st.rerun()


def get_student(student_id: str):
    if not student_id:
        return None
    return next((s for s in st.session_state.students if s["id"] == student_id), None)


def get_active_staff():
    return st.session_state.staff


def get_session_from_time(t: time) -> str:
    h = t.hour
    if h < 11:
        return "Morning"
    elif h < 13:
        return "Middle"
    else:
        return "Afternoon"


def calculate_age(dob_str: str):
    try:
        d = datetime.strptime(dob_str, "%Y-%m-%d").date()
        today = date.today()
        years = today.year - d.year - ((today.month, today.day) < (d.month, d.day))
        return years
    except Exception:
        return None


def generate_mock_incidents(n: int = 70):
    """Create random quick incidents so the analysis page has something to show."""
    incidents = []
    
    # Weight certain students to have more incidents for realistic patterns
    student_weights = {
        "stu_jp1": 8,
        "stu_jp2": 5, 
        "stu_jp3": 3,
        "stu_py1": 10,
        "stu_py2": 7,
        "stu_py3": 4,
        "stu_sy1": 12,
        "stu_sy2": 9,
        "stu_sy3": 6,
    }
    
    student_pool = []
    for stu in MOCK_STUDENTS:
        weight = student_weights.get(stu["id"], 5)
        student_pool.extend([stu] * weight)
    
    for _ in range(n):
        stu = random.choice(student_pool)
        beh = random.choice(BEHAVIOUR_TYPES)
        ant = random.choice(ANTECEDENTS)
        loc = random.choice(LOCATIONS)
        support = random.choice(SUPPORT_TYPES)
        interv = random.choice(INTERVENTIONS)
        
        # More realistic severity distribution
        sev = random.choices([1, 2, 3, 4, 5], weights=[20, 35, 25, 15, 5])[0]

        dt = datetime.now() - timedelta(days=random.randint(0, 90))
        t_hour = random.choices([9, 10, 11, 12, 13, 14, 15], weights=[10, 15, 12, 8, 12, 18, 10])[0]
        dt = dt.replace(hour=t_hour, minute=random.randint(0, 59), second=0)

        incidents.append(
            {
                "id": str(uuid.uuid4()),
                "student_id": stu["id"],
                "student_name": stu["name"],
                "date": dt.date().isoformat(),
                "time": dt.time().strftime("%H:%M:%S"),
                "day": dt.strftime("%A"),
                "session": get_session_from_time(dt.time()),
                "location": loc,
                "behaviour_type": beh,
                "antecedent": ant,
                "support_type": support,
                "intervention": interv,
                "severity": sev,
                "reported_by": random.choice(MOCK_STAFF)["name"],
                "additional_staff": [],
                "description": "Auto-generated mock incident.",
                "hypothesis": generate_simple_function(ant, beh),
                "is_critical": sev >= 4,
                "duration_minutes": random.randint(2, 25),
            }
        )
    return incidents


# =========================================
# HYPOTHESIS ENGINE (SIMPLE)
# =========================================


def generate_simple_function(antecedent: str, behaviour: str) -> str:
    """
    Returns one simple hypothesis line in the requested format:
    'To get/avoid <function>' where <function> is one of:
    tangible / request / activity / sensory / attention.
    """
    ant = (antecedent or "").lower()
    beh = (behaviour or "").lower()

    # Decide get vs avoid and function
    if any(k in ant for k in ["instruction", "demand", "work", "task", "academic"]):
        get_avoid = "To avoid"
        fn = "request / activity"
    elif "transition" in ant:
        get_avoid = "To avoid"
        fn = "transition / activity"
    elif any(k in ant for k in ["denied", "access", "item", "object", "preferred"]):
        get_avoid = "To get"
        fn = "tangible"
    elif any(k in ant for k in ["sensory", "noise", "lights", "overload"]):
        get_avoid = "To avoid"
        fn = "sensory"
    elif any(k in ant for k in ["peer", "attention", "staff"]):
        get_avoid = "To get"
        fn = "attention"
    else:
        get_avoid = "To get"
        fn = "attention"

    return f"{get_avoid} {fn}"


# =========================================
# WORD DOCUMENT GENERATION WITH GRAPHS
# =========================================


def save_plotly_as_image(fig, width=6, height=4):
    """Save a plotly figure as an image and return the path"""
    try:
        # Create a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        temp_path = temp_file.name
        temp_file.close()
        
        # Save the figure
        fig.write_image(temp_path, width=width*100, height=height*100, scale=2)
        return temp_path
    except Exception as e:
        st.warning(f"Could not save graph: {e}")
        return None


def generate_behaviour_analysis_plan_docx(student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level):
    """Generate a Word document for Behaviour Analysis Plan with embedded graphs"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Behaviour Analysis Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Student Info
    doc.add_heading('Student Information', 1)
    student_table = doc.add_table(rows=4, cols=2)
    student_table.style = 'Light Grid Accent 1'
    
    student_table.rows[0].cells[0].text = 'Student Name:'
    student_table.rows[0].cells[1].text = student['name']
    student_table.rows[1].cells[0].text = 'Program:'
    student_table.rows[1].cells[1].text = student['program']
    student_table.rows[2].cells[0].text = 'Grade:'
    student_table.rows[2].cells[1].text = student['grade']
    student_table.rows[3].cells[0].text = 'Report Date:'
    student_table.rows[3].cells[1].text = datetime.now().strftime('%d/%m/%Y')
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Total Incidents: ").bold = True
    summary_para.add_run(f"{len(full_df)}\n")
    summary_para.add_run(f"Critical Incidents: ").bold = True
    summary_para.add_run(f"{len(full_df[full_df['incident_type'] == 'Critical'])}\n")
    summary_para.add_run(f"Average Severity: ").bold = True
    summary_para.add_run(f"{full_df['severity'].mean():.2f}\n")
    summary_para.add_run(f"Risk Level: ").bold = True
    summary_para.add_run(f"{risk_level} ({risk_score}/100)")
    
    doc.add_paragraph()
    
    # ===== GRAPH 1: Incident Frequency Over Time =====
    doc.add_heading('Incident Patterns Over Time', 1)
    
    try:
        daily_counts = full_df.groupby(full_df["date_parsed"].dt.date).size().reset_index(name="count")
        fig1 = go.Figure()
        fig1.add_trace(go.Scatter(
            x=daily_counts["date_parsed"],
            y=daily_counts["count"],
            mode='lines+markers',
            name='Incidents per day',
            line=dict(color='#2c5282', width=2),
            fill='tozeroy',
            fillcolor='rgba(44, 82, 130, 0.2)'
        ))
        fig1.update_layout(
            title="Daily Incident Frequency",
            xaxis_title="Date",
            yaxis_title="Number of Incidents",
            template="plotly_white",
            height=350
        )
        
        img_path = save_plotly_as_image(fig1, width=6, height=3.5)
        if img_path:
            doc.add_picture(img_path, width=Inches(6))
            os.unlink(img_path)
    except Exception as e:
        doc.add_paragraph(f"[Graph unavailable: {str(e)}]")
    
    doc.add_paragraph()
    
    # ===== GRAPH 2: Day-of-Week Heatmap =====
    doc.add_heading('High-Risk Times Analysis', 1)
    
    try:
        full_df["hour"] = pd.to_datetime(full_df["time"], format="%H:%M:%S", errors="coerce").dt.hour
        full_df["day_of_week"] = full_df["date_parsed"].dt.day_name()
        
        day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
        pivot = full_df.pivot_table(
            values="severity",
            index="day_of_week",
            columns="hour",
            aggfunc="count",
            fill_value=0
        )
        pivot = pivot.reindex([d for d in day_order if d in pivot.index], fill_value=0)
        
        fig2 = go.Figure(data=go.Heatmap(
            z=pivot.values,
            x=pivot.columns,
            y=pivot.index,
            colorscale='Blues',
            showscale=True
        ))
        fig2.update_layout(
            title="Incident Frequency by Day & Hour",
            xaxis_title="Hour of Day",
            yaxis_title="Day of Week",
            template="plotly_white",
            height=350
        )
        
        img_path = save_plotly_as_image(fig2, width=6, height=3.5)
        if img_path:
            doc.add_picture(img_path, width=Inches(6))
            os.unlink(img_path)
    except Exception as e:
        doc.add_paragraph(f"[Graph unavailable: {str(e)}]")
    
    doc.add_paragraph()
    
    # ===== GRAPH 3: Behaviour Type Distribution =====
    doc.add_heading('Behaviour Type Analysis', 1)
    
    try:
        beh_counts = full_df["behaviour_type"].value_counts()
        
        fig3 = go.Figure(data=[go.Bar(
            x=beh_counts.index,
            y=beh_counts.values,
            marker=dict(color='#2c5282')
        )])
        fig3.update_layout(
            title="Distribution of Behaviour Types",
            xaxis_title="Behaviour Type",
            yaxis_title="Frequency",
            template="plotly_white",
            height=350
        )
        fig3.update_xaxes(tickangle=-45)
        
        img_path = save_plotly_as_image(fig3, width=6, height=3.5)
        if img_path:
            doc.add_picture(img_path, width=Inches(6))
            os.unlink(img_path)
    except Exception as e:
        doc.add_paragraph(f"[Graph unavailable: {str(e)}]")
    
    doc.add_paragraph()
    
    # Data Findings
    doc.add_heading('Summary of Data Findings', 1)
    findings = doc.add_paragraph()
    findings.add_run('Primary Concern: ').bold = True
    findings.add_run(f"{top_beh} is the most frequently recorded behaviour of concern.\n\n")
    findings.add_run('Key Triggers: ').bold = True
    findings.add_run(f"The most common antecedent is '{top_ant}', indicating this context regularly precedes dysregulation.\n\n")
    findings.add_run('Hotspot Locations: ').bold = True
    findings.add_run(f"Incidents most often occur in {top_loc}, particularly during the {top_session} session.")
    
    doc.add_paragraph()
    
    # Clinical Interpretation
    doc.add_heading('Clinical Interpretation (Trauma-Informed)', 1)
    clinical = doc.add_paragraph()
    clinical.add_run(
        f"Patterns suggest that {student['name']} is most vulnerable when '{top_ant}' occurs, "
        f"often in the {top_loc} during {top_session}. These moments likely narrow the student's "
        "window of tolerance, increasing the risk of fight/flight responses.\n\n"
    )
    clinical.add_run(
        "Through a trauma-informed lens, this behaviour is understood as a safety strategy rather than "
        "wilful defiance. CPI emphasises staying in the Supportive phase as early as possible.\n\n"
    )
    clinical.add_run(
        "The Berry Street Education Model points towards strengthening Body (regulation routines, "
        "predictable transitions) and Relationship (connection before correction)."
    )
    
    doc.add_paragraph()
    
    # Recommendations
    doc.add_heading('Recommendations & Next Steps', 1)
    
    doc.add_heading('1. Proactive Regulation Around Key Triggers', 2)
    rec1 = doc.add_paragraph(style='List Bullet')
    rec1.add_run(f"Provide a brief check-in and clear visual cue before '{top_ant}'")
    rec2 = doc.add_paragraph(style='List Bullet')
    rec2.add_run(f"Offer a regulated start before the high-risk {top_session} session")
    
    doc.add_heading('2. Co-regulation & Staff Responses (CPI Aligned)', 2)
    rec3 = doc.add_paragraph(style='List Bullet')
    rec3.add_run("Use CPI Supportive stance, low slow voice and minimal language")
    rec4 = doc.add_paragraph(style='List Bullet')
    rec4.add_run("Reduce audience by moving peers where possible")
    
    doc.add_heading('3. Teaching Replacement Skills', 2)
    rec5 = doc.add_paragraph(style='List Bullet')
    rec5.add_run("Link goals to Personal and Social Capability")
    rec6 = doc.add_paragraph(style='List Bullet')
    rec6.add_run("Explicitly teach help-seeking routines")
    
    doc.add_heading('4. SMART Goal Example', 2)
    goal = doc.add_paragraph(style='List Bullet')
    goal.add_run(
        f"Over the next 5 weeks, during identified trigger times, the student will use "
        f"an agreed help-seeking strategy in 4 out of 5 opportunities."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('\n\nGenerated by CLC Behaviour Support System\n')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_para.add_run(datetime.now().strftime('%d %B %Y at %I:%M %p'))
    
    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# =========================================
# PAGES
# =========================================


def render_login_page():
    st.markdown("## 🔐 Staff Login")
    st.caption("Enter your email and password to access the system")
    
    with st.expander("📧 Demo Credentials"):
        st.markdown("**Staff Accounts:**")
        for staff in MOCK_STAFF[:3]:
            st.code(f"Email: {staff['email']}\nPassword: demo123")
        st.markdown("**Admin Account:**")
        st.code(f"Email: admin.user@example.com\nPassword: admin123")
        st.info("💡 You can also use 'demo' as a universal password for any staff email")

    email = st.text_input("Email address", placeholder="emily.jones@example.com")
    password = st.text_input("Password", type="password", placeholder="Enter password")

    if st.button("Login", type="primary"):
        if not email or not password:
            st.error("⚠️ Please enter both email and password")
        else:
            if login_user(email, password):
                st.success(f"✅ Logged in as {st.session_state.current_user['name']}")
                st.rerun()
            else:
                st.error("❌ Invalid email or password. Please try again.")


def render_landing_page():
    user = st.session_state.current_user or {}
    st.markdown(
        f"### 👋 Welcome, **{user.get('name', 'User')}** "
        f"({user.get('role', 'Role unknown')})"
    )
    st.caption(f"Email: {user.get('email', 'N/A')}")

    col1, col2 = st.columns([4, 1])
    with col2:
        if st.button("Logout", use_container_width=True):
            st.session_state.logged_in = False
            st.session_state.current_user = None
            st.session_state.current_page = "login"
            st.rerun()

    st.markdown("---")

    # Quick stats
    total_incidents = len(st.session_state.incidents)
    total_students = len(st.session_state.students)
    critical_count = len([i for i in st.session_state.incidents if i.get("is_critical")])
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("📚 Total Students", total_students)
    with col2:
        st.metric("📊 Total Incidents", total_incidents)
    with col3:
        st.metric("🚨 Critical Incidents", critical_count)

    st.markdown("---")
    st.markdown("### 📚 Select Program")

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("#### Junior Primary")
        jp_students = len([s for s in st.session_state.students if s["program"] == "JP"])
        st.caption(f"{jp_students} students")
        if st.button("Enter JP", key="enter_jp", use_container_width=True, type="primary"):
            go_to("program_students", selected_program="JP")
            
    with col2:
        st.markdown("#### Primary Years")
        py_students = len([s for s in st.session_state.students if s["program"] == "PY"])
        st.caption(f"{py_students} students")
        if st.button("Enter PY", key="enter_py", use_container_width=True, type="primary"):
            go_to("program_students", selected_program="PY")
            
    with col3:
        st.markdown("#### Senior Years")
        sy_students = len([s for s in st.session_state.students if s["program"] == "SY"])
        st.caption(f"{sy_students} students")
        if st.button("Enter SY", key="enter_sy", use_container_width=True, type="primary"):
            go_to("program_students", selected_program="SY")

    st.markdown("---")

    st.markdown("### ⚡ Quick Actions")
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### Quick Incident Log")
        all_students = [s for s in st.session_state.students]
        selected = st.selectbox(
            "Select student",
            options=all_students,
            format_func=lambda s: f"{s['name']} ({s['program']} - Grade {s['grade']})",
        )

        if st.button("Start Quick Log", type="primary", key="quick_log_btn", use_container_width=True):
            go_to("incident_log", selected_student_id=selected["id"])
    
    with col2:
        st.markdown("#### Program Overview")
        st.caption("View cross-program analytics")
        if st.button("View Program Analytics", use_container_width=True):
            go_to("program_overview")


def render_program_students_page():
    program = st.session_state.get("selected_program", "JP")
    st.markdown(f"## {PROGRAM_NAMES.get(program, program)} Program — Students")

    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("⬅ Back to landing"):
            go_to("landing")

    students = [s for s in st.session_state.students if s["program"] == program]

    if not students:
        st.info("No students in this program (mock).")
        return

    # Student cards with incident counts
    for stu in students:
        stu_incidents = [i for i in st.session_state.incidents if i["student_id"] == stu["id"]]
        critical = [i for i in stu_incidents if i.get("is_critical")]
        
        with st.container(border=True):
            col1, col2, col3 = st.columns([3, 2, 2])
            
            with col1:
                st.markdown(f"### {stu['name']}")
                st.caption(f"Grade {stu['grade']} | EDID {stu['edid']}")
                age = calculate_age(stu["dob"])
                if age is not None:
                    st.caption(f"Age: {age} years")
            
            with col2:
                st.metric("Incidents", len(stu_incidents))
                st.caption(f"Critical: {len(critical)}")
            
            with col3:
                if st.button("📝 Log Incident", key=f"log_{stu['id']}", use_container_width=True):
                    go_to("incident_log", selected_student_id=stu["id"])
                if st.button("📊 Analysis", key=f"ana_{stu['id']}", use_container_width=True):
                    go_to("student_analysis", selected_student_id=stu["id"])


def render_incident_log_page():
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected.")
        if st.button("Back to landing"):
            go_to("landing")
        return

    st.markdown(f"## 📝 Quick Incident Log — {student['name']}")
    st.caption(f"{student['program']} Program | Grade {student['grade']}")

    with st.form("incident_form"):
        col1, col2 = st.columns(2)
        with col1:
            inc_date = st.date_input("Date", date.today())
            inc_time = st.time_input("Time", datetime.now().time())
            location = st.selectbox("Location", LOCATIONS)
        with col2:
            behaviour = st.selectbox("Behaviour Type", BEHAVIOUR_TYPES)
            antecedent = st.selectbox("Antecedent (Trigger)", ANTECEDENTS)
            support_type = st.selectbox("Type of Support", SUPPORT_TYPES)

        st.markdown("### Staff")
        reporter = st.session_state.current_user["name"]
        st.info(f"Staff member reporting: **{reporter}** (auto-filled)")

        staff_names = [s["name"] for s in get_active_staff()]
        additional_staff = st.multiselect(
            "Additional staff involved (if any)", options=staff_names
        )

        st.markdown("### Intervention Used")
        intervention = st.selectbox(
            "Adult action / de-escalation strategy",
            INTERVENTIONS,
        )
        
        duration = st.number_input("Duration (minutes)", min_value=1, max_value=60, value=5)

        severity = st.slider("Severity (1 = low, 5 = high)", 1, 5, 2)

        description = st.text_area(
            "Brief description (factual)",
            placeholder="Short, objective description of what occurred…",
        )

        st.markdown("### Hypothesis (auto-generated, editable)")
        auto_hyp = generate_simple_function(antecedent, behaviour)
        hypothesis = st.text_input("Function of behaviour", value=auto_hyp)

        submitted = st.form_submit_button("Submit Incident", type="primary")

    if submitted:
        new_id = str(uuid.uuid4())
        rec = {
            "id": new_id,
            "student_id": student_id,
            "student_name": student["name"],
            "date": inc_date.isoformat(),
            "time": inc_time.strftime("%H:%M:%S"),
            "day": inc_date.strftime("%A"),
            "session": get_session_from_time(inc_time),
            "location": location,
            "behaviour_type": behaviour,
            "antecedent": antecedent,
            "support_type": support_type,
            "reported_by": reporter,
            "additional_staff": additional_staff,
            "intervention": intervention,
            "severity": severity,
            "duration_minutes": duration,
            "description": description,
            "hypothesis": hypothesis,
            "is_critical": severity >= 4,
        }
        st.session_state.incidents.append(rec)
        st.success("✅ Incident saved (sandbox).")

        if severity >= 4:
            st.warning("⚠️ Severity ≥ 4 → Proceed to Critical Incident ABCH form?")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📋 Complete ABCH Form", type="primary"):
                    go_to("critical_incident", current_incident_id=new_id)
            with col2:
                if st.button("↩️ Back to students"):
                    go_to("program_students", selected_program=student["program"])
        else:
            if st.button("↩️ Back to students"):
                go_to("program_students", selected_program=student["program"])


def generate_critical_recommendations(quick_inc: dict) -> str:
    ant = quick_inc.get("antecedent", "known triggers")
    beh = quick_inc.get("behaviour_type", "behaviour of concern")
    loc = quick_inc.get("location", "learning area")
    sess = quick_inc.get("session", "session")

    return (
        f"Patterns suggest incidents often occur when '{ant}' in the {loc} during {sess}. "
        "Consider applying CPI's Supportive stance earlier in the escalation, increasing "
        "co-regulation opportunities, and building predictable routines at these times. "
        "Use Berry Street Education Model strategies (Body and Relationship) such as grounding, "
        "rhythmic regulation, and warm relational check-ins. Link adjustments to the Australian "
        "Curriculum General Capabilities (Personal & Social Capability), and set a SMART goal "
        "around help-seeking and self-regulation during similar tasks or transitions."
    )


def render_critical_incident_page():
    inc_id = st.session_state.get("current_incident_id")
    quick_inc = next((i for i in st.session_state.incidents if i["id"] == inc_id), None)

    if not quick_inc:
        st.error("No quick incident found to build the critical form from.")
        if st.button("Back to landing"):
            go_to("landing")
        return

    st.markdown("## 🚨 Critical Incident ABCH Form")
    st.caption("Auto-filled from quick log. Edit as required.")

    with st.expander("Quick incident details"):
        st.json(quick_inc)

    # ABCH laid out in 4 columns
    st.markdown("### ABCH Overview")
    colA, colB, colC, colH = st.columns(4)

    with colA:
        st.subheader("A – Antecedent")
        A_text = st.text_area(
            "What happened before?",
            value=quick_inc.get("antecedent", ""),
            key="crit_A",
        )

    with colB:
        st.subheader("B – Behaviour")
        B_text = st.text_area(
            "What did the student do?",
            value=quick_inc.get("behaviour_type", ""),
            key="crit_B",
        )

    with colC:
        st.subheader("C – Consequence")
        C_text = st.text_area(
            "What happened after?",
            value="",
            key="crit_C",
        )

    with colH:
        st.subheader("H – Hypothesis")
        default_H = generate_simple_function(
            quick_inc.get("antecedent", ""), quick_inc.get("behaviour_type", "")
        )
        H_text = st.text_area("Why did this occur?", value=default_H, key="crit_H")

    st.markdown("---")

    # Additional ABCH lines for complex incidents
    st.markdown("### Additional Incident Elements (optional)")

    if st.button("➕ Add another ABCH line"):
        st.session_state.abch_rows.append({"A": "", "B": "", "C": "", "H": ""})

    for idx, row in enumerate(st.session_state.abch_rows):
        st.markdown(f"**Extra element {idx+1}**")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            row["A"] = st.text_input("A", value=row["A"], key=f"rowA_{idx}")
        with c2:
            row["B"] = st.text_input("B", value=row["B"], key=f"rowB_{idx}")
        with c3:
            row["C"] = st.text_input("C", value=row["C"], key=f"rowC_{idx}")
        with c4:
            row["H"] = st.text_input("H", value=row["H"], key=f"rowH_{idx}")

    st.markdown("---")

    st.markdown("### Safety Responses (CPI-aligned, non-restraint)")
    safety_responses = st.multiselect(
        "Actions taken",
        [
            "CPI Supportive stance",
            "Cleared nearby students",
            "Student moved to safer location",
            "Additional staff attended",
            "Safety plan enacted",
            "Continued monitoring until regulated",
            "First aid offered",
        ],
    )

    st.markdown("### Notifications")
    notifications = st.multiselect(
        "Who was notified?",
        [
            "Parent / carer",
            "Line manager",
            "Safety & Wellbeing / SSS",
            "DCP",
            "SAPOL",
            "First Aid officer",
            "Injury report completed",
            "Transport home required",
        ],
    )

    st.markdown("### Outcome actions")
    removed = st.checkbox("Removed from learning")
    family_contact = st.checkbox("Family contacted")
    safety_updated = st.checkbox("Safety plan updated")
    transport_home = st.checkbox("Transport home required")
    other_actions = st.text_area("Other actions / follow-up")

    st.markdown("### Recommendations (auto-generated, editable)")
    rec_text = generate_critical_recommendations(quick_inc)
    recommendations = st.text_area("Recommendations", value=rec_text, height=200)

    if st.button("Save critical incident", type="primary"):
        record = {
            "id": str(uuid.uuid4()),
            "created_at": datetime.now().isoformat(),
            "quick_incident_id": quick_inc["id"],
            "student_id": quick_inc["student_id"],
            "ABCH_primary": {"A": A_text, "B": B_text, "C": C_text, "H": H_text},
            "ABCH_additional": st.session_state.abch_rows,
            "safety_responses": safety_responses,
            "notifications": notifications,
            "outcomes": {
                "removed": removed,
                "family_contact": family_contact,
                "safety_updated": safety_updated,
                "transport_home": transport_home,
                "other": other_actions,
            },
            "recommendations": recommendations,
        }
        st.session_state.critical_incidents.append(record)
        st.success("✅ Critical incident saved (sandbox).")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("📊 Go to student analysis"):
                go_to("student_analysis", selected_student_id=quick_inc["student_id"])
        with col2:
            if st.button("↩️ Back to students"):
                go_to("program_students", selected_program=get_student(quick_inc["student_id"])["program"])


# =========================================
# ADVANCED STUDENT ANALYSIS PAGE
# =========================================

def render_student_analysis_page():
    """
    Comprehensive analytics dashboard with 20+ visualizations
    """
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected.")
        if st.button("Back to landing"):
            go_to("landing")
        return

    st.markdown(f"## 📊 Advanced Data Analysis — {student['name']}")
    st.caption(f"{student['program']} program | Grade {student['grade']}")

    # Get incidents
    quick = [i for i in st.session_state.incidents if i["student_id"] == student_id]
    crit = [c for c in st.session_state.critical_incidents if c["student_id"] == student_id]

    if not quick and not crit:
        st.info("No incident data yet for this student.")
        if st.button("Log first incident"):
            go_to("incident_log", selected_student_id=student_id)
        return

    # Build unified dataframe
    quick_df = pd.DataFrame(quick) if quick else pd.DataFrame()
    crit_df = pd.DataFrame(crit) if crit else pd.DataFrame()

    if not quick_df.empty:
        quick_df["incident_type"] = "Quick"
        quick_df["date_parsed"] = pd.to_datetime(quick_df["date"])

    if not crit_df.empty:
        crit_df["incident_type"] = "Critical"
        if "created_at" in crit_df.columns:
            crit_df["date_parsed"] = pd.to_datetime(crit_df["created_at"])
        else:
            crit_df["date_parsed"] = pd.to_datetime(datetime.now().isoformat())
        crit_df["severity"] = 5
        crit_df["antecedent"] = crit_df["ABCH_primary"].apply(
            lambda d: d.get("A") if isinstance(d, dict) else ""
        )
        crit_df["behaviour_type"] = crit_df["ABCH_primary"].apply(
            lambda d: d.get("B") if isinstance(d, dict) else ""
        )

    full_df = pd.concat([quick_df, crit_df], ignore_index=True)
    full_df = full_df.sort_values("date_parsed")

    # ==============================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ==============================================
    st.markdown("## 📈 Executive Summary")
    
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.metric("Total Incidents", len(full_df))
    with col2:
        critical_count = len(full_df[full_df["incident_type"] == "Critical"])
        st.metric("Critical", critical_count)
    with col3:
        avg_sev = round(full_df["severity"].mean(), 2)
        st.metric("Avg Severity", avg_sev)
    with col4:
        days_span = (full_df["date_parsed"].max() - full_df["date_parsed"].min()).days + 1
        st.metric("Days Tracked", days_span)
    with col5:
        incidents_per_day = round(len(full_df) / days_span, 2)
        st.metric("Inc/Day", incidents_per_day)

    # Trend indicator
    if len(full_df) >= 2:
        recent_avg = full_df.tail(5)["severity"].mean()
        older_avg = full_df.head(5)["severity"].mean()
        trend = "📈 Increasing" if recent_avg > older_avg else "📉 Decreasing" if recent_avg < older_avg else "➡️ Stable"
        st.info(f"**Severity Trend (last 5 vs first 5):** {trend}")

    st.markdown("---")

    # ==============================================
    # SECTION 2: TIME-SERIES ANALYSIS
    # ==============================================
    st.markdown("## ⏰ Time-Series Analysis")

    # 2.1 Daily incident frequency
    st.markdown("### 📅 Incident Frequency Over Time")
    daily_counts = full_df.groupby(full_df["date_parsed"].dt.date).size().reset_index(name="count")
    fig1 = go.Figure()
    fig1.add_trace(go.Scatter(
        x=daily_counts["date_parsed"],
        y=daily_counts["count"],
        mode='lines+markers',
        name='Incidents per day',
        line=dict(color='#2c5282', width=2),
        fill='tozeroy',
        fillcolor='rgba(44, 82, 130, 0.2)'
    ))
    fig1.update_layout(
        title="Daily Incident Count",
        xaxis_title="Date",
        yaxis_title="Number of Incidents",
        hovermode='x unified',
        template="plotly_white"
    )
    st.plotly_chart(fig1, use_container_width=True)

    # 2.2 Moving average (7-day)
    if len(full_df) >= 7:
        st.markdown("### 📊 7-Day Moving Average (Smoothed Trend)")
        full_df["date_only"] = full_df["date_parsed"].dt.date
        daily = full_df.groupby("date_only").size().reset_index(name="count")
        daily["7d_avg"] = daily["count"].rolling(window=7, min_periods=1).mean()
        
        fig2 = go.Figure()
        fig2.add_trace(go.Scatter(
            x=daily["date_only"],
            y=daily["count"],
            mode='lines',
            name='Daily count',
            line=dict(color='lightgray', width=1)
        ))
        fig2.add_trace(go.Scatter(
            x=daily["date_only"],
            y=daily["7d_avg"],
            mode='lines',
            name='7-day average',
            line=dict(color='#dc2626', width=3)
        ))
        fig2.update_layout(
            title="Trend Analysis (7-Day Moving Average)",
            xaxis_title="Date",
            yaxis_title="Incidents",
            template="plotly_white"
        )
        st.plotly_chart(fig2, use_container_width=True)

    # 2.3 Severity timeline with annotations
    st.markdown("### 🎯 Severity Timeline (with Critical Incidents Highlighted)")
    fig3 = go.Figure()
    
    # Regular incidents
    quick_only = full_df[full_df["incident_type"] == "Quick"]
    crit_only = full_df[full_df["incident_type"] == "Critical"]
    
    if not quick_only.empty:
        fig3.add_trace(go.Scatter(
            x=quick_only["date_parsed"],
            y=quick_only["severity"],
            mode='markers',
            name='Quick Incident',
            marker=dict(size=10, color='#2c5282', opacity=0.6),
            hovertemplate='%{y} - %{text}<extra></extra>',
            text=quick_only["behaviour_type"]
        ))
    
    if not crit_only.empty:
        fig3.add_trace(go.Scatter(
            x=crit_only["date_parsed"],
            y=crit_only["severity"],
            mode='markers',
            name='Critical Incident',
            marker=dict(size=15, color='#dc2626', symbol='star'),
            hovertemplate='CRITICAL: %{text}<extra></extra>',
            text=crit_only["behaviour_type"]
        ))
    
    fig3.update_layout(
        title="Severity Over Time (Quick vs Critical)",
        xaxis_title="Date",
        yaxis_title="Severity Level",
        yaxis=dict(range=[0, 6]),
        template="plotly_white"
    )
    st.plotly_chart(fig3, use_container_width=True)

    st.markdown("---")

    # ==============================================
    # SECTION 3: HEATMAPS & PATTERN ANALYSIS
    # ==============================================
    st.markdown("## 🔥 Heatmaps & Pattern Analysis")

    # 3.1 Day of week vs Time of day heatmap
    st.markdown("### 🗓️ Day-of-Week × Time-of-Day Heatmap")
    full_df["hour"] = pd.to_datetime(full_df["time"], format="%H:%M:%S", errors="coerce").dt.hour
    full_df["day_of_week"] = full_df["date_parsed"].dt.day_name()
    
    # Create pivot table
    day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    pivot = full_df.pivot_table(
        values="severity",
        index="day_of_week",
        columns="hour",
        aggfunc="count",
        fill_value=0
    )
    pivot = pivot.reindex(day_order, fill_value=0)
    
    fig4 = go.Figure(data=go.Heatmap(
        z=pivot.values,
        x=pivot.columns,
        y=pivot.index,
        colorscale='Blues',
        hovertemplate='%{y}, %{x}:00<br>Incidents: %{z}<extra></extra>'
    ))
    fig4.update_layout(
        title="Incident Frequency by Day & Hour",
        xaxis_title="Hour of Day",
        yaxis_title="Day of Week",
        template="plotly_white"
    )
    st.plotly_chart(fig4, use_container_width=True)

    # 3.2 Location vs Session heatmap
    st.markdown("### 📍 Location × Session Heatmap")
    loc_sess_pivot = full_df.pivot_table(
        values="severity",
        index="location",
        columns="session",
        aggfunc="count",
        fill_value=0
    )
    
    fig5 = go.Figure(data=go.Heatmap(
        z=loc_sess_pivot.values,
        x=loc_sess_pivot.columns,
        y=loc_sess_pivot.index,
        colorscale='Teal',
        hovertemplate='%{y} during %{x}<br>Incidents: %{z}<extra></extra>'
    ))
    fig5.update_layout(
        title="Location Hotspots by Session",
        xaxis_title="Session",
        yaxis_title="Location",
        template="plotly_white"
    )
    st.plotly_chart(fig5, use_container_width=True)

    st.markdown("---")

    # ==============================================
    # SECTION 4: BEHAVIOUR PATTERN ANALYSIS
    # ==============================================
    st.markdown("## 🧩 Behaviour Pattern Analysis")

    # 4.1 Antecedent-Behaviour co-occurrence
    st.markdown("### 🔗 Antecedent → Behaviour Patterns")
    
    ant_beh_counts = full_df.groupby(["antecedent", "behaviour_type"]).size().reset_index(name="count")
    ant_beh_counts = ant_beh_counts.sort_values("count", ascending=False).head(15)
    
    fig6 = go.Figure(data=[go.Bar(
        x=ant_beh_counts["count"],
        y=[f"{row['antecedent'][:30]}... → {row['behaviour_type']}" 
           for _, row in ant_beh_counts.iterrows()],
        orientation='h',
        marker=dict(
            color=ant_beh_counts["count"],
            colorscale='Teal'
        )
    )])
    fig6.update_layout(
        title="Top 15 Antecedent → Behaviour Pairs",
        xaxis_title="Frequency",
        yaxis_title="Pattern",
        template="plotly_white"
    )
    st.plotly_chart(fig6, use_container_width=True)

    # 4.2 Behaviour type distribution (pie chart)
    st.markdown("### 🥧 Behaviour Type Distribution")
    beh_counts = full_df["behaviour_type"].value_counts()
    
    fig7 = go.Figure(data=[go.Pie(
        labels=beh_counts.index,
        values=beh_counts.values,
        hole=0.3,
        marker=dict(colors=['#2c5282', '#0d9488', '#059669', '#2563eb', '#7c3aed', '#db2777'])
    )])
    fig7.update_layout(
        title="Behaviour Type Breakdown",
        template="plotly_white"
    )
    st.plotly_chart(fig7, use_container_width=True)

    # 4.3 Behaviour chains (sequences)
    st.markdown("### 🔄 Behaviour Sequences (What Follows What)")
    if len(full_df) >= 3:
        full_df_sorted = full_df.sort_values("date_parsed")
        sequences = []
        for i in range(len(full_df_sorted) - 1):
            curr = full_df_sorted.iloc[i]["behaviour_type"]
            next_beh = full_df_sorted.iloc[i + 1]["behaviour_type"]
            sequences.append(f"{curr} → {next_beh}")
        
        seq_counts = pd.Series(sequences).value_counts().head(10)
        
        fig8 = go.Figure(data=[go.Bar(
            x=seq_counts.values,
            y=seq_counts.index,
            orientation='h',
            marker=dict(color='#7c3aed')
        )])
        fig8.update_layout(
            title="Top 10 Behaviour Sequences",
            xaxis_title="Frequency",
            yaxis_title="Sequence",
            template="plotly_white"
        )
        st.plotly_chart(fig8, use_container_width=True)

    st.markdown("---")

    # ==============================================
    # SECTION 5: INTERVENTION EFFECTIVENESS
    # ==============================================
    st.markdown("## 🎯 Intervention Effectiveness Analysis")

    # 5.1 Intervention vs Severity
    st.markdown("### 📊 Intervention Success Rate (Severity Reduction)")
    
    interv_sev = full_df.groupby("intervention").agg({
        "severity": ["mean", "count"]
    }).reset_index()
    interv_sev.columns = ["intervention", "avg_severity", "count"]
    interv_sev = interv_sev[interv_sev["count"] >= 2]
    interv_sev = interv_sev.sort_values("avg_severity")
    
    fig9 = go.Figure()
    fig9.add_trace(go.Bar(
        x=interv_sev["avg_severity"],
        y=interv_sev["intervention"],
        orientation='h',
        marker=dict(
            color=interv_sev["avg_severity"],
            colorscale='RdYlGn',
            reversescale=True,
            cmin=1,
            cmax=5
        ),
        text=interv_sev["count"],
        texttemplate='n=%{text}',
        textposition='outside'
    ))
    fig9.update_layout(
        title="Average Severity by Intervention Type (Lower = More Effective)",
        xaxis_title="Average Severity",
        yaxis_title="Intervention",
        xaxis=dict(range=[0, 5.5]),
        template="plotly_white"
    )
    st.plotly_chart(fig9, use_container_width=True)

    # 5.2 Duration analysis
    if "duration_minutes" in full_df.columns:
        st.markdown("### ⏱️ Incident Duration Analysis")
        
        fig10 = go.Figure()
        fig10.add_trace(go.Box(
            y=full_df["duration_minutes"],
            x=full_df["behaviour_type"],
            marker=dict(color='#2c5282')
        ))
        fig10.update_layout(
            title="Duration Distribution by Behaviour Type",
            xaxis_title="Behaviour Type",
            yaxis_title="Duration (minutes)",
            template="plotly_white"
        )
        fig10.update_xaxes(tickangle=-45)
        st.plotly_chart(fig10, use_container_width=True)

    st.markdown("---")

    # ==============================================
    # SECTION 6: PREDICTIVE INDICATORS
    # ==============================================
    st.markdown("## 🔮 Predictive Indicators & Risk Analysis")

    # 6.1 Escalation pattern detection
    st.markdown("### ⚠️ Escalation Pattern Detection")
    
    full_df["severity_change"] = full_df["severity"].diff()
    escalations = full_df[full_df["severity_change"] > 0]
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Escalation Events", len(escalations))
        st.caption("Times when severity increased from previous incident")
    with col2:
        if len(escalations) > 0:
            avg_escalation = escalations["severity_change"].mean()
            st.metric("Avg Escalation Jump", f"+{avg_escalation:.1f}")
        else:
            st.metric("Avg Escalation Jump", "N/A")

    # 6.2 Risk score calculation
    st.markdown("### 🎲 Current Risk Assessment")
    
    recent_incidents = full_df.tail(5)
    risk_factors = {
        "Recent frequency": len(full_df.tail(7)) / 7,
        "Recent avg severity": recent_incidents["severity"].mean(),
        "Critical incident rate": (len(full_df[full_df["incident_type"] == "Critical"]) / len(full_df)) * 100,
        "Escalation trend": 1 if len(full_df) >= 2 and full_df.tail(5)["severity"].mean() > full_df.head(5)["severity"].mean() else 0
    }
    
    risk_score = min(100, int(
        (risk_factors["Recent frequency"] * 10) +
        (risk_factors["Recent avg severity"] * 8) +
        (risk_factors["Critical incident rate"] * 0.5) +
        (risk_factors["Escalation trend"] * 20)
    ))
    
    risk_color = "#059669" if risk_score < 30 else "#d97706" if risk_score < 60 else "#dc2626"
    risk_level = "LOW" if risk_score < 30 else "MODERATE" if risk_score < 60 else "HIGH"
    
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        st.markdown(f"### Overall Risk Score: <span style='color:{risk_color}; font-size:2em;'>{risk_score}/100</span>", 
                   unsafe_allow_html=True)
    with col2:
        st.markdown(f"### Level: <span style='color:{risk_color};'>{risk_level}</span>", 
                   unsafe_allow_html=True)
    
    with st.expander("📊 Risk Factor Breakdown"):
        for factor, value in risk_factors.items():
            st.metric(factor, f"{value:.2f}")

    st.markdown("---")

    # ==============================================
    # SECTION 7: ABC HYPOTHESIS ANALYSIS
    # ==============================================
    st.markdown("## 🧠 Functional Behaviour Analysis")

    if "hypothesis" in full_df.columns:
        st.markdown("### 🎯 Hypothesized Functions (ABC Analysis)")
        
        functions = full_df["hypothesis"].value_counts()
        
        fig12 = go.Figure(data=[go.Bar(
            x=functions.values,
            y=functions.index,
            orientation='h',
            marker=dict(
                color=['#dc2626', '#d97706', '#059669', '#2c5282', '#7c3aed'][:len(functions)]
            )
        )])
        fig12.update_layout(
            title="Behavioural Function Distribution",
            xaxis_title="Frequency",
            yaxis_title="Function",
            template="plotly_white"
        )
        st.plotly_chart(fig12, use_container_width=True)

        top_function = functions.index[0]
        st.info(f"**Primary Function:** {top_function} ({functions.values[0]} incidents, "
               f"{(functions.values[0]/len(full_df)*100):.1f}% of total)")

    st.markdown("---")

    # ==============================================
    # SECTION 8: CLINICAL INTERPRETATION
    # ==============================================
    if not full_df.empty:
        top_ant = full_df["antecedent"].mode()[0] if len(full_df["antecedent"]) > 0 else "Unknown"
        top_beh = full_df["behaviour_type"].mode()[0] if len(full_df["behaviour_type"]) > 0 else "Unknown"
        top_loc = full_df["location"].mode()[0] if len(full_df["location"]) > 0 else "Unknown"
        top_session = full_df["session"].mode()[0] if len(full_df["session"]) > 0 else "Unknown"

        total = len(full_df)
        crit_total = len(full_df[full_df["incident_type"] == "Critical"])
        quick_total = total - crit_total
        crit_rate = (crit_total / total) * 100 if total > 0 else 0

        full_sorted = full_df.sort_values("date_parsed")
        if len(full_sorted) >= 2:
            first_sev = full_sorted["severity"].iloc[0]
            last_sev = full_sorted["severity"].iloc[-1]
            if last_sev > first_sev:
                severity_trend = "increasing over time"
            elif last_sev < first_sev:
                severity_trend = "decreasing over time"
            else:
                severity_trend = "relatively stable over time"
        else:
            severity_trend = "unable to determine (limited data)"

        st.markdown("## 🧠 Clinical Interpretation & Next Steps")

        st.markdown("### 1. Summary of Data Findings")

        st.markdown(
            f"- **Primary concern:** **{top_beh}** is the most frequently recorded behaviour of concern."
        )
        st.markdown(
            f"- **Key triggers:** The most common antecedent is **{top_ant}**, "
            f"indicating this context regularly precedes dysregulation."
        )
        st.markdown(
            f"- **Hotspot locations:** Incidents most often occur in **{top_loc}**, "
            f"particularly during the **{top_session}** session."
        )
        st.markdown(
            f"- **Incident profile:** {quick_total} quick incidents and {crit_total} "
            f"critical incidents have been recorded (critical incidents = "
            f"**{crit_rate:.1f}%** of all incidents)."
        )
        st.markdown(
            f"- **Severity trend:** Overall severity appears **{severity_trend}**."
        )

        st.markdown("### 2. Clinical Interpretation (Trauma-Informed)")

        clinical_text = (
            f"Patterns suggest that {student['name']} is most vulnerable when **{top_ant}** "
            f"occurs, often in the **{top_loc}** during **{top_session}**. These moments "
            "likely narrow the student's window of tolerance, increasing the risk of "
            "fight/flight responses such as the identified behaviour.\n\n"
            "Through a **trauma-informed lens**, this behaviour is understood as a safety "
            "strategy rather than wilful defiance. CPI emphasises staying in the **Supportive** "
            "phase as early as possible — calm body language, non-threatening stance and "
            "minimal verbal load.\n\n"
            "The **Berry Street Education Model** (Body, Relationship, Stamina, Engagement) "
            "points towards strengthening **Body** (regulation routines, predictable transitions) "
            "and **Relationship** (connection before correction). SMART trauma principles "
            "highlight the importance of predictability, relational safety and reducing cognitive "
            "load during known trigger times."
        )
        st.info(clinical_text)

        st.markdown("### 3. Next Steps & Recommendations")

        next_steps = (
            "1. **Proactive regulation around key triggers**  \n"
            f"   - Provide a brief check-in and clear visual cue before **{top_ant}**.  \n"
            "   - Offer a regulated start (breathing, movement, sensory tool) before the "
            f"high-risk **{top_session}** session.\n\n"
            "2. **Co-regulation & staff responses (CPI aligned)**  \n"
            "   - Use CPI Supportive stance, low slow voice and minimal language when early "
            "signs of escalation appear.  \n"
            "   - Reduce audience by moving peers where possible and maintain connection with "
            "one key adult.\n\n"
            "3. **Teaching replacement skills (Australian Curriculum – General Capabilities)**  \n"
            "   - Link goals to **Personal and Social Capability** (self-management & "
            "social management).  \n"
            "   - Explicitly teach and rehearse a help-seeking routine the student can use "
            "in place of the behaviour (e.g., card, phrase, movement to a safe space).\n\n"
            "4. **SMART-style goal example**  \n"
            "   - *Over the next 5 weeks, during identified trigger times, the student will "
            "use an agreed help-seeking strategy instead of the behaviour of concern in "
            "4 out of 5 opportunities, with co-regulation support from staff.*"
        )
        st.success(next_steps)

    # ==============================================
    # SECTION 9: DATA EXPORT
    # ==============================================
    st.markdown("---")
    st.markdown("## 📄 Data Export & Reporting")

    col1, col2 = st.columns(2)
    
    with col1:
        csv = full_df.to_csv(index=False)
        st.download_button(
            label="📥 Download Full Dataset (CSV)",
            data=csv,
            file_name=f"{student['name']}_incidents.csv",
            mime="text/csv"
        )
    
    with col2:
        # Generate Word document with graphs
        with st.spinner("Generating report with graphs... This may take a moment."):
            try:
                docx_file = generate_behaviour_analysis_plan_docx(
                    student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level
                )
                
                if docx_file:
                    st.download_button(
                        label="📄 Download Behaviour Analysis Plan (Word + Graphs)",
                        data=docx_file,
                        file_name=f"Behaviour_Analysis_Plan_{student['name'].replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"Error generating document: {e}")
                st.info("Note: Graph embedding requires kaleido package. Install with: pip install kaleido")

    st.markdown("---")
    if st.button("⬅ Back to Students", type="primary"):
        go_to("program_students", selected_program=student["program"])


# NEW: Program Overview Page
def render_program_overview_page():
    st.markdown("## 📈 Cross-Program Analytics")
    st.caption("Incident patterns across all programs")
    
    if st.button("⬅ Back to landing"):
        go_to("landing")
    
    incidents = st.session_state.incidents
    if not incidents:
        st.info("No incidents recorded yet.")
        return
    
    df = pd.DataFrame(incidents)
    df["date_parsed"] = pd.to_datetime(df["date"])
    
    df["program"] = df["student_id"].apply(
        lambda sid: next((s["program"] for s in st.session_state.students if s["id"] == sid), "Unknown")
    )
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Incidents", len(df))
    with col2:
        st.metric("Critical Incidents", len(df[df["is_critical"] == True]))
    with col3:
        st.metric("Average Severity", round(df["severity"].mean(), 2))
    
    st.markdown("---")
    
    st.markdown("### 📚 Incidents by Program")
    prog_counts = df["program"].value_counts().reset_index()
    prog_counts.columns = ["Program", "Count"]
    fig1 = px.bar(prog_counts, x="Program", y="Count", color="Program", 
                  color_discrete_map={"JP": "#2c5282", "PY": "#0d9488", "SY": "#7c3aed"},
                  template="plotly_white")
    st.plotly_chart(fig1, use_container_width=True)
    
    st.markdown("### ⚠️ Behaviour Types by Program")
    beh_by_prog = df.groupby(["program", "behaviour_type"]).size().reset_index(name="count")
    fig2 = px.bar(beh_by_prog, x="behaviour_type", y="count", color="program", barmode="group",
                  color_discrete_map={"JP": "#2c5282", "PY": "#0d9488", "SY": "#7c3aed"},
                  template="plotly_white")
    fig2.update_xaxes(tickangle=-45)
    st.plotly_chart(fig2, use_container_width=True)
    
    st.markdown("### 📅 Incident Trends Over Time")
    df["week"] = df["date_parsed"].dt.to_period("W").astype(str)
    weekly = df.groupby(["week", "program"]).size().reset_index(name="count")
    fig3 = px.line(weekly, x="week", y="count", color="program",
                   color_discrete_map={"JP": "#2c5282", "PY": "#0d9488", "SY": "#7c3aed"},
                   template="plotly_white")
    fig3.update_xaxes(tickangle=-45)
    st.plotly_chart(fig3, use_container_width=True)
    
    st.markdown("### 📍 Location Hotspots (All Programs)")
    loc_counts = df["location"].value_counts().head(10).reset_index()
    loc_counts.columns = ["Location", "Count"]
    fig4 = px.bar(loc_counts, x="Count", y="Location", orientation="h", 
                  color_discrete_sequence=["#2c5282"],
                  template="plotly_white")
    st.plotly_chart(fig4, use_container_width=True)


# =========================================
# MAIN APP ROUTER
# =========================================

def main():
    init_state()

    if not st.session_state.logged_in:
        render_login_page()
        return

    page = st.session_state.current_page

    if page == "landing":
        render_landing_page()
    elif page == "program_students":
        render_program_students_page()
    elif page == "incident_log":
        render_incident_log_page()
    elif page == "critical_incident":
        render_critical_incident_page()
    elif page == "student_analysis":
        render_student_analysis_page()
    elif page == "program_overview":
        render_program_overview_page()
    elif page == "login":
        render_login_page()
    else:
        st.error("Unknown page.")
        render_landing_page()


if __name__ == "__main__":
    main()
